import csv

from datetime import date, timedelta
from io import StringIO
from io import BytesIO
from pathlib import Path



from flask import Blueprint, Response, render_template



from models.repositories import query_all, query_one

from routes.audit import log_action

from routes.utils import permission_required



exports_bp = Blueprint("exports", __name__, url_prefix="/exports")


@exports_bp.route("/")
@permission_required("exports.view")
def reports_home():
    return render_template("reports.html")





def _csv_response(filename: str, rows, headers):

    output = StringIO()

    writer = csv.writer(output)

    writer.writerow(headers)

    writer.writerows(rows)

    return Response(

        output.getvalue(),

        mimetype="text/csv",

        headers={"Content-Disposition": f"attachment; filename={filename}"},

    )


def _to_lines(prefix: str, text: str, max_chars: int = 90):
    clean = (text or "").strip() or "-"
    lines = []
    current = prefix
    for word in clean.split():
        if len(current) + len(word) + 1 <= max_chars:
            current = f"{current} {word}".strip()
        else:
            lines.append(current)
            current = f"    {word}"
    lines.append(current)
    return lines


def _weekly_report_context():
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    week_end = week_start + timedelta(days=4)
    week_start_s = week_start.isoformat()
    week_end_s = week_end.isoformat()

    # Shown in footer (center) only, not in header/cover.
    sales_contact = "sales@ishango-it.com"
    row_email = query_one("SELECT value FROM app_settings WHERE key = 'sales_contact_email'")
    if row_email and row_email["value"]:
        sales_contact = row_email["value"].strip()
    author_name = "Josué Mbuyu wa Kabinga"
    row_author = query_one("SELECT value FROM app_settings WHERE key = 'report_author_name'")
    if row_author and row_author["value"]:
        author_name = row_author["value"].strip()
    author_email = "josue.mbuyu@ishango-it.com"
    row_author_email = query_one("SELECT value FROM app_settings WHERE key = 'report_author_email'")
    if row_author_email and row_author_email["value"]:
        author_email = row_author_email["value"].strip()
    author_title = "Technical sales"
    row_author_title = query_one("SELECT value FROM app_settings WHERE key = 'report_author_title'")
    if row_author_title and row_author_title["value"]:
        author_title = row_author_title["value"].strip()
    app_name = "Sales Manager"
    row_name = query_one("SELECT value FROM app_settings WHERE key = 'app_name'")
    if row_name and row_name["value"]:
        app_name = row_name["value"].strip()

    project_updates = query_all(
        """
        SELECT p.title AS project_title, pu.summary, pu.next_action, pu.next_action_date, pu.risk_level, pu.is_blocked, pu.created_at
        FROM project_updates pu
        JOIN projects p ON p.id = pu.project_id
        WHERE date(pu.created_at) BETWEEN date(?) AND date(?)
        ORDER BY pu.created_at DESC
        LIMIT 25
        """,
        (week_start_s, week_end_s),
    )
    delivered = query_all(
        """
        SELECT title, progress, due_date
        FROM projects
        WHERE progress >= 100
          AND date(due_date) BETWEEN date(?) AND date(?)
        ORDER BY due_date ASC
        LIMIT 20
        """,
        (week_start_s, week_end_s),
    )
    blockers = query_all(
        """
        SELECT p.title, pu.summary, pu.created_at
        FROM project_updates pu
        JOIN projects p ON p.id = pu.project_id
        WHERE pu.is_blocked = 1
          AND date(pu.created_at) BETWEEN date(?) AND date(?)
        ORDER BY pu.created_at DESC
        LIMIT 20
        """,
        (week_start_s, week_end_s),
    )
    next_priorities = query_all(
        """
        SELECT p.title, pu.next_action, pu.next_action_date, pu.risk_level
        FROM project_updates pu
        JOIN projects p ON p.id = pu.project_id
        WHERE pu.next_action IS NOT NULL
          AND pu.next_action <> ''
          AND pu.next_action_date IS NOT NULL
          AND pu.next_action_date <> ''
          AND date(pu.next_action_date) BETWEEN date(?) AND date(?, '+7 day')
        ORDER BY date(pu.next_action_date) ASC
        LIMIT 20
        """,
        (week_start_s, week_end_s),
    )

    # Central tracking (Suivis)
    tracking_stats = query_one(
        """
        SELECT
          COUNT(*) AS total,
          SUM(CASE WHEN status = 'pending' THEN 1 ELSE 0 END) AS pending,
          SUM(CASE WHEN status = 'done' THEN 1 ELSE 0 END) AS done,
          SUM(CASE WHEN status = 'pending'
                    AND due_date IS NOT NULL AND due_date <> ''
                    AND date(due_date) < date('now')
              THEN 1 ELSE 0 END) AS overdue
        FROM tracking_items
        WHERE date(created_at) BETWEEN date(?) AND date(?)
        """,
        (week_start_s, week_end_s),
    )
    tracking_stats = dict(tracking_stats) if tracking_stats else {}
    tracking_urgent = query_all(
        """
        SELECT
          title, context_label, due_date, priority, source_type, is_auto
        FROM tracking_items
        WHERE status = 'pending'
          AND due_date IS NOT NULL AND due_date <> ''
          AND date(due_date) BETWEEN date(?) AND date(?, '+14 day')
        ORDER BY
          CASE priority
            WHEN 'critique' THEN 1
            WHEN 'haute' THEN 2
            WHEN 'moyenne' THEN 3
            ELSE 4
          END,
          date(due_date) ASC,
          updated_at DESC
        LIMIT 12
        """,
        (week_start_s, week_end_s),
    )

    # Appointments
    appointment_stats = query_one(
        """
        SELECT
          COUNT(*) AS total,
          SUM(CASE WHEN status = 'pending' THEN 1 ELSE 0 END) AS pending,
          SUM(CASE WHEN status = 'done' THEN 1 ELSE 0 END) AS done,
          SUM(CASE WHEN status = 'pending' AND date(date) < date('now') THEN 1 ELSE 0 END) AS overdue
        FROM appointments
        WHERE date(date) BETWEEN date(?) AND date(?)
        """,
        (week_start_s, week_end_s),
    )
    appointment_stats = dict(appointment_stats) if appointment_stats else {}
    appointment_updates = query_all(
        """
        SELECT
          c.name AS client_name,
          au.summary,
          au.status_snapshot,
          au.date_snapshot,
          au.time_snapshot,
          au.created_at
        FROM appointment_updates au
        JOIN appointments a ON a.id = au.appointment_id
        JOIN clients c ON c.id = a.client_id
        WHERE date(au.created_at) BETWEEN date(?) AND date(?)
        ORDER BY au.created_at DESC
        LIMIT 15
        """,
        (week_start_s, week_end_s),
    )
    appointment_overdue = query_all(
        """
        SELECT c.name AS client_name, a.date, a.time, a.location
        FROM appointments a
        JOIN clients c ON c.id = a.client_id
        WHERE a.status = 'pending'
          AND date(a.date) < date('now')
        ORDER BY date(a.date) ASC, a.time ASC
        LIMIT 10
        """
    )

    # Supplier deliveries (risks)
    supplier_risk_stats = query_one(
        """
        SELECT
          SUM(CASE WHEN d.status NOT IN ('delivered','validated')
                    AND COALESCE(d.expected_date, d.planned_date) IS NOT NULL
                    AND COALESCE(d.expected_date, d.planned_date) <> ''
                    AND date(COALESCE(d.expected_date, d.planned_date)) < date('now')
              THEN 1 ELSE 0 END) AS late_count,
          SUM(CASE WHEN d.status NOT IN ('delivered','validated') AND d.blocker IS NOT NULL AND d.blocker <> '' THEN 1 ELSE 0 END) AS blocked_count,
          SUM(CASE WHEN d.status NOT IN ('delivered','validated') AND d.quality_status = 'issue' THEN 1 ELSE 0 END) AS quality_issue_count,
          SUM(CASE WHEN d.status NOT IN ('delivered','validated') AND ((
                    COALESCE(d.expected_date, d.planned_date) IS NOT NULL
                    AND COALESCE(d.expected_date, d.planned_date) <> ''
                    AND date(COALESCE(d.expected_date, d.planned_date)) < date('now')
                  ) OR (d.blocker IS NOT NULL AND d.blocker <> '') OR d.quality_status = 'issue')
              THEN COALESCE(d.amount, 0) ELSE 0 END) AS risk_value
        FROM supplier_deliveries d
        WHERE date(d.created_at) BETWEEN date(?) AND date(?)
        """,
        (week_start_s, week_end_s),
    )
    supplier_risk_stats = dict(supplier_risk_stats) if supplier_risk_stats else {}
    supplier_risk_items = query_all(
        """
        SELECT
          s.name AS supplier_name,
          d.title,
          COALESCE(d.expected_date, d.planned_date) AS due_date,
          d.status,
          d.blocker,
          d.quality_status,
          CASE
            WHEN COALESCE(d.expected_date, d.planned_date) IS NOT NULL
                 AND COALESCE(d.expected_date, d.planned_date) <> ''
                 AND date(COALESCE(d.expected_date, d.planned_date)) < date('now')
            THEN CAST(julianday('now') - julianday(COALESCE(d.expected_date, d.planned_date)) AS INT)
            ELSE 0
          END AS days_late
        FROM supplier_deliveries d
        JOIN suppliers s ON s.id = d.supplier_id
        WHERE d.status NOT IN ('delivered','validated')
          AND (
            (COALESCE(d.expected_date, d.planned_date) IS NOT NULL AND COALESCE(d.expected_date, d.planned_date) <> '' AND date(COALESCE(d.expected_date, d.planned_date)) <= date(?, '+7 day'))
            OR (d.blocker IS NOT NULL AND d.blocker <> '')
            OR d.quality_status = 'issue'
          )
        ORDER BY
          CASE
            WHEN d.blocker IS NOT NULL AND d.blocker <> '' THEN 1
            WHEN d.quality_status = 'issue' THEN 2
            WHEN days_late > 0 THEN 3
            ELSE 4
          END,
          days_late DESC,
          date(COALESCE(d.expected_date, d.planned_date)) ASC
        LIMIT 12
        """,
        (week_end_s,),
    )
    return {
        "week_start_s": week_start_s,
        "week_end_s": week_end_s,
        "sales_contact": sales_contact,
        "author_name": author_name,
        "author_email": author_email,
        "author_title": author_title,
        "app_name": app_name,
        "project_updates": project_updates,
        "delivered": delivered,
        "blockers": blockers,
        "next_priorities": next_priorities,
        "tracking_stats": tracking_stats,
        "tracking_urgent": tracking_urgent,
        "appointment_stats": appointment_stats,
        "appointment_updates": appointment_updates,
        "appointment_overdue": appointment_overdue,
        "supplier_risk_stats": supplier_risk_stats,
        "supplier_risk_items": supplier_risk_items,
    }


def _build_weekly_report_pdf_bytes(ctx):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    week_start_s = ctx["week_start_s"]
    week_end_s = ctx["week_end_s"]
    sales_contact = ctx["sales_contact"]
    author_name = ctx.get("author_name") or ""
    author_email = ctx.get("author_email") or ""
    author_title = ctx.get("author_title") or ""
    app_name = ctx["app_name"]
    project_updates = ctx["project_updates"]
    delivered = ctx["delivered"]
    blockers = ctx["blockers"]
    next_priorities = ctx["next_priorities"]
    tracking_stats = ctx.get("tracking_stats") or {}
    tracking_urgent = ctx.get("tracking_urgent") or []
    appointment_stats = ctx.get("appointment_stats") or {}
    appointment_updates = ctx.get("appointment_updates") or []
    appointment_overdue = ctx.get("appointment_overdue") or []
    supplier_risk_stats = ctx.get("supplier_risk_stats") or {}
    supplier_risk_items = ctx.get("supplier_risk_items") or []

    stream = BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    # Layout constants (consistent margins/spacing)
    M = 44  # side margin
    HEADER_H = 82
    FOOTER_LINE_Y = 52
    FOOTER_TEXT_Y = 34
    CONTENT_TOP_Y = height - 130
    CONTENT_MIN_Y = FOOTER_LINE_Y + 18
    CONTENT_W = width - (M * 2)

    y = CONTENT_TOP_Y
    brand_orange = colors.HexColor("#F68B1F")
    brand_blue = colors.HexColor("#1E3A8A")
    text_dark = colors.HexColor("#1F2937")
    muted = colors.HexColor("#6B7280")
    logo_path = Path(__file__).resolve().parent.parent / "static" / "img" / "ishango-logo.png"

    def ensure_space(needed: int = 22):
        nonlocal y
        if y < CONTENT_MIN_Y + needed:
            pdf.showPage()
            draw_page_header()
            draw_page_footer()
            y = CONTENT_TOP_Y

    def draw_page_header():
        pdf.setFillColor(colors.HexColor("#F9FAFB"))
        pdf.rect(0, height - HEADER_H, width, HEADER_H, stroke=0, fill=1)
        pdf.setStrokeColor(colors.HexColor("#E5E7EB"))
        pdf.line(M, height - HEADER_H, width - M, height - HEADER_H)
        if logo_path.exists():
            try:
                pdf.drawImage(str(logo_path), M, height - 68, width=78, height=34, mask="auto")
            except Exception:
                pass
        pdf.setFillColor(brand_blue)
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(M + 90, height - 44, f"{app_name} - Rapport hebdomadaire")
        pdf.setFillColor(muted)
        pdf.setFont("Helvetica", 9)
        pdf.drawString(M + 90, height - 59, f"Période : du {week_start_s} au {week_end_s}")
        # contact email moved to footer (center)

    def draw_page_footer():
        pdf.setStrokeColor(colors.HexColor("#E5E7EB"))
        pdf.line(M, FOOTER_LINE_Y, width - M, FOOTER_LINE_Y)
        pdf.setFillColor(muted)
        pdf.setFont("Helvetica", 9)
        pdf.drawString(M, FOOTER_TEXT_Y, "ISHANGO IT Solutions SARL")
        if sales_contact:
            pdf.drawCentredString(width / 2, FOOTER_TEXT_Y, sales_contact)
        pdf.drawRightString(width - M, FOOTER_TEXT_Y, f"Édition : {date.today().isoformat()}")

    def section_title(text: str):
        nonlocal y
        # Keep consistent vertical centering across pages/sections.
        ensure_space(32)
        bar_h = 22
        bar_y = y - bar_h
        pdf.setFillColor(brand_blue)
        pdf.roundRect(M, bar_y, CONTENT_W, bar_h, 6, stroke=0, fill=1)
        pdf.setFillColor(colors.white)
        pdf.setFont("Helvetica-Bold", 11)
        # Baseline positioning for better centering
        text_y = bar_y + 7
        pdf.drawString(M + 10, text_y, text)
        y -= (bar_h + 14)

    def _wrap_text(text: str, max_width: float, font_name: str, font_size: float):
        words = (text or "").strip().split()
        if not words:
            return ["-"]
        lines = []
        cur = ""
        for w in words:
            test = (cur + " " + w).strip()
            if pdf.stringWidth(test, font_name, font_size) <= max_width:
                cur = test
            else:
                if cur:
                    lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    def bullet(text: str, color=text_dark):
        nonlocal y
        font_name = "Helvetica"
        font_size = 10
        bullet_symbol = "•"
        gap = 10
        indent = 14
        max_w = CONTENT_W - (indent + gap + 8)
        lines = _wrap_text(text, max_w, font_name, font_size)
        needed = 14 * len(lines) + 4
        ensure_space(needed)
        pdf.setFillColor(color)
        pdf.setFont(font_name, font_size)
        # first line with bullet symbol
        pdf.drawString(M + 8, y, bullet_symbol)
        pdf.drawString(M + 8 + indent, y, lines[0])
        y -= 14
        # wrapped lines aligned with text start
        for line in lines[1:]:
            pdf.drawString(M + 8 + indent, y, line)
            y -= 14

    def draw_cover_page():
        pdf.setFillColor(colors.white)
        pdf.rect(0, 0, width, height, stroke=0, fill=1)
        pdf.setFillColor(brand_orange)
        pdf.rect(0, height - 120, width, 120, stroke=0, fill=1)
        if logo_path.exists():
            try:
                pdf.drawImage(str(logo_path), M, height - 94, width=120, height=52, mask="auto")
            except Exception:
                pass
        pdf.setFillColor(colors.white)
        pdf.setFont("Helvetica-Bold", 26)
        pdf.drawString(M + 130, height - 66, "Rapport Hebdomadaire")
        pdf.setFillColor(text_dark)
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(M, height - 170, app_name)
        pdf.setFont("Helvetica", 12)
        pdf.drawString(M, height - 194, f"Période : du {week_start_s} au {week_end_s}")
        # Do not show sales contact on cover; keep it in footer only
        if author_name or author_email or author_title:
            pdf.setFillColor(muted)
            pdf.setFont("Helvetica", 11)
            if author_name:
                pdf.drawString(M, height - 236, f"Fait par : {author_name}")
            if author_title:
                pdf.drawString(M, height - 254, f"Poste : {author_title}")
            if author_email:
                pdf.drawString(M, height - 272, f"Email : {author_email}")

        box_y = height - 520
        pdf.setFillColor(colors.HexColor("#F9FAFB"))
        pdf.setStrokeColor(colors.HexColor("#E5E7EB"))
        pdf.roundRect(M, box_y, CONTENT_W, 180, 10, stroke=1, fill=1)
        pdf.setFillColor(brand_blue)
        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawString(M + 14, box_y + 152, "Résumé exécutif")
        # KPI layout (modern, aligned, with emphasis on numbers)
        kpi_x = M + 14
        kpi_w = CONTENT_W - 28
        row_h = 22
        start_y = box_y + 126

        def kpi_row(idx: int, label: str, value: str, *, tone: str = "neutral"):
            y_row = start_y - (idx * row_h)
            # subtle alternating background
            if idx % 2 == 0:
                pdf.setFillColor(colors.white)
                pdf.roundRect(kpi_x, y_row - 6, kpi_w, row_h, 6, stroke=0, fill=1)

            pdf.setFillColor(text_dark)
            pdf.setFont("Helvetica", 10.5)
            pdf.drawString(kpi_x + 10, y_row + 2, label)

            if tone == "danger":
                pill_bg = colors.HexColor("#DC2626")
                pill_fg = colors.white
            elif tone == "warning":
                pill_bg = colors.HexColor("#F59E0B")
                pill_fg = colors.white
            elif tone == "info":
                pill_bg = colors.HexColor("#1E3A8A")
                pill_fg = colors.white
            else:
                pill_bg = colors.HexColor("#111827")
                pill_fg = colors.white

            # value pill on the right
            pill_w = max(64, min(140, 8 * len(value) + 26))
            pill_h = 16
            pill_x = kpi_x + kpi_w - pill_w - 10
            pill_y = y_row
            pdf.setFillColor(pill_bg)
            pdf.roundRect(pill_x, pill_y, pill_w, pill_h, 8, stroke=0, fill=1)
            pdf.setFillColor(pill_fg)
            pdf.setFont("Helvetica-Bold", 10.5)
            pdf.drawCentredString(pill_x + (pill_w / 2), pill_y + 4.5, value)

        blocked_n = len(blockers)
        delivered_n = len(delivered)
        priorities_n = len(next_priorities)
        pending_n = int(tracking_stats.get("pending") or 0)
        overdue_n = int(tracking_stats.get("overdue") or 0)
        ap_total_n = int(appointment_stats.get("total") or 0)
        ap_done_n = int(appointment_stats.get("done") or 0)

        kpi_row(0, "Activités enregistrées", str(len(project_updates)), tone="info")
        kpi_row(1, "Blocages détectés", str(blocked_n), tone="danger" if blocked_n > 0 else "neutral")
        kpi_row(2, "Livrables terminés", str(delivered_n), tone="neutral")
        kpi_row(3, "Priorités (7 prochains jours)", str(priorities_n), tone="warning" if priorities_n > 0 else "neutral")
        kpi_row(4, "Suivis (en attente / en retard)", f"{pending_n} / {overdue_n}", tone="warning" if overdue_n > 0 else "neutral")
        kpi_row(5, "Rendez-vous (planifiés / terminés)", f"{ap_total_n} / {ap_done_n}", tone="neutral")

        # Use the same footer as all other pages (left brand + centered email + edition).
        draw_page_footer()
        pdf.showPage()

    draw_cover_page()
    draw_page_header()
    draw_page_footer()

    card_y = height - 114
    card_w = (CONTENT_W - 20) / 3
    cards = [
        ("Activités", str(len(project_updates))),
        ("Blocages", str(len(blockers))),
        ("Livrables", str(len(delivered))),
    ]
    for idx, (label, value) in enumerate(cards):
        x = M + idx * (card_w + 10)
        pdf.setFillColor(colors.white)
        pdf.setStrokeColor(colors.HexColor("#D1D5DB"))
        pdf.roundRect(x, card_y, card_w, 34, 6, stroke=1, fill=1)
        pdf.setFillColor(muted)
        pdf.setFont("Helvetica", 9)
        pdf.drawString(x + 8, card_y + 20, label)
        pdf.setFillColor(text_dark)
        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawRightString(x + card_w - 8, card_y + 18, value)

    section_title("1) Avancement de la semaine")
    if project_updates:
        for item in project_updates[:10]:
            # Two-line compact bullet: "Projet: résumé"
            prefix = f"{item['project_title']}:"
            summary = (item["summary"] or "").strip()
            combined = f"{prefix} {summary}".strip()
            bullet(combined)
            flags = []
            if item["is_blocked"]:
                flags.append("BLOQUÉ")
            if item["risk_level"] == "eleve":
                flags.append("RISQUE ÉLEVÉ")
            if flags:
                bullet(f"[{', '.join(flags)}]", color=brand_orange)
    else:
        bullet("Aucun avancement saisi cette semaine.")
    y -= 6

    section_title("2) Blocages à lever")
    if blockers:
        for item in blockers[:8]:
            for txt in _to_lines(f"{item['title']}:", item["summary"], max_chars=95):
                bullet(txt)
    else:
        bullet("Aucun blocage sur la période.")
    y -= 6

    section_title("3) Livrables terminés")
    if delivered:
        for item in delivered[:8]:
            bullet(f"{item['title']} (100%)")
    else:
        bullet("Aucun livrable finalisé sur la période.")
    y -= 6

    section_title("4) Priorités (7 prochains jours)")
    if next_priorities:
        for item in next_priorities[:10]:
            bullet(
                f"{item['title']} | {item['next_action']} | Échéance : {item['next_action_date']} | Risque : {item['risk_level']}"
            )
    else:
        bullet("Aucune priorité planifiée.")
    y -= 6

    section_title("5) Suivis centralisés (opérations)")
    pending = tracking_stats.get("pending") or 0
    overdue = tracking_stats.get("overdue") or 0
    done = tracking_stats.get("done") or 0
    bullet(f"Suivis en attente : {pending} | En retard : {overdue} | Terminés : {done}")
    if tracking_urgent:
        for item in tracking_urgent[:10]:
            due = item["due_date"] or "-"
            ctx_label = (item["context_label"] or "").strip()
            pr = item["priority"] or "moyenne"
            label = f"{item['title']} | Échéance : {due} | Priorité : {pr}"
            if ctx_label:
                label = f"{label} | {ctx_label}"
            bullet(label)
    else:
        bullet("Aucun suivi prioritaire sur la période.")
    y -= 6

    section_title("6) Rendez-vous (activité & retards)")
    ap_total = appointment_stats.get("total") or 0
    ap_pending = appointment_stats.get("pending") or 0
    ap_done = appointment_stats.get("done") or 0
    ap_overdue = appointment_stats.get("overdue") or 0
    bullet(
        f"Planifiés (semaine) : {ap_total} | En attente : {ap_pending} | Terminés : {ap_done} | En retard : {ap_overdue}"
    )
    if appointment_updates:
        for item in appointment_updates[:8]:
            who = (item["client_name"] or "-").strip()
            sumy = (item["summary"] or "-").strip()
            snap_d = item["date_snapshot"] or "-"
            snap_t = item["time_snapshot"] or "-"
            bullet(f"{who} : {sumy} | {snap_d} {snap_t}")
    else:
        bullet("Aucune mise à jour de rendez-vous sur la période.")
    if appointment_overdue:
        y -= 2
        bullet("Principaux rendez-vous en retard :", color=brand_orange)
        for item in appointment_overdue[:5]:
            bullet(f"{item['client_name']} | {item['date']} {item['time']} | {item['location'] or '-'}")
    y -= 6

    section_title("7) Fournisseurs / livraisons (risques)")
    late_c = supplier_risk_stats.get("late_count") or 0
    blk_c = supplier_risk_stats.get("blocked_count") or 0
    qual_c = supplier_risk_stats.get("quality_issue_count") or 0
    risk_v = supplier_risk_stats.get("risk_value") or 0
    bullet(f"Retards : {late_c} | Blocages : {blk_c} | Non-conformités : {qual_c} | Valeur à risque : {risk_v}")
    if supplier_risk_items:
        for item in supplier_risk_items[:10]:
            flags = []
            if (item["blocker"] or "").strip():
                flags.append("BLOCAGE")
            if item["quality_status"] == "issue":
                flags.append("QUALITÉ")
            if (item["days_late"] or 0) > 0:
                flags.append(f"RETARD +{item['days_late']}j")
            flag_s = f"[{', '.join(flags)}] " if flags else ""
            bullet(f"{flag_s}{item['supplier_name']} | {item['title']} | Échéance : {item['due_date'] or '-'}")
    else:
        bullet("Aucun risque fournisseur prioritaire.")
    y -= 6

    section_title("8) Points de décision (manager)")
    bullet("Arbitrer les blocages critiques (projets, fournisseurs) et valider les délais repoussés.")
    bullet("Prioriser les suivis en retard et les actions à forte valeur client.")

    pdf.save()
    stream.seek(0)
    filename = f"rapport_hebdo_{week_start_s}_au_{week_end_s}.pdf"
    return stream.getvalue(), filename





@exports_bp.route("/clients.csv")

@permission_required("exports.view")

def export_clients():

    rows = query_all(

        "SELECT name, company, pipeline_stage, phone, email, notes, created_at FROM clients ORDER BY name ASC"

    )

    log_action("export", "clients", "csv")

    return _csv_response(

        "clients.csv",

        [tuple(row) for row in rows],

        ["Nom", "Entreprise", "Pipeline", "Telephone", "Email", "Notes", "Cree le"],

    )





@exports_bp.route("/appointments.csv")

@permission_required("exports.view")

def export_appointments():

    rows = query_all(

        """

        SELECT c.name, a.date, a.time, a.location, a.status, a.notes

        FROM appointments a

        JOIN clients c ON c.id = a.client_id

        ORDER BY a.date ASC, a.time ASC

        """

    )

    log_action("export", "appointments", "csv")

    return _csv_response(

        "appointments.csv",

        [tuple(row) for row in rows],

        ["Client", "Date", "Heure", "Lieu", "Statut", "Notes"],

    )


@exports_bp.route("/weekly-report.pdf")
@permission_required("exports.view")
def export_weekly_report_pdf():
    try:
        import reportlab  # noqa: F401
    except Exception:
        return Response(
            "La dependance 'reportlab' est requise pour exporter le PDF. Installez-la puis relancez l'application.",
            mimetype="text/plain",
            status=500,
        )
    ctx = _weekly_report_context()
    pdf_bytes, filename = _build_weekly_report_pdf_bytes(ctx)
    log_action("export", "weekly_report", "pdf", {"from": ctx["week_start_s"], "to": ctx["week_end_s"]})
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@exports_bp.route("/weekly-report/preview")
@permission_required("exports.view")
def preview_weekly_report_pdf():
    try:
        import reportlab  # noqa: F401
    except Exception:
        return Response(
            "La dependance 'reportlab' est requise pour exporter le PDF. Installez-la puis relancez l'application.",
            mimetype="text/plain",
            status=500,
        )
    ctx = _weekly_report_context()
    pdf_bytes, filename = _build_weekly_report_pdf_bytes(ctx)
    log_action("preview", "weekly_report", "pdf", {"from": ctx["week_start_s"], "to": ctx["week_end_s"]})
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f"inline; filename={filename}"},
    )


@exports_bp.route("/weekly-report.docx")
@permission_required("exports.view")
def export_weekly_report_docx():
    try:
        from docx import Document
    except Exception:
        return Response(
            "La dependance 'python-docx' est requise pour exporter Word. Installez-la puis relancez l'application.",
            mimetype="text/plain",
            status=500,
        )
    ctx = _weekly_report_context()
    doc = Document()
    doc.add_heading("Rapport hebdomadaire", 0)
    doc.add_paragraph(f"Periode: {ctx['week_start_s']} au {ctx['week_end_s']}")
    doc.add_paragraph(ctx["sales_contact"])

    doc.add_heading("1) Avancement de la semaine", level=1)
    if ctx["project_updates"]:
        for item in ctx["project_updates"][:10]:
            doc.add_paragraph(f"{item['project_title']}: {item['summary']}", style="List Bullet")
    else:
        doc.add_paragraph("Aucun avancement saisi cette semaine.", style="List Bullet")

    doc.add_heading("2) Blocages a lever", level=1)
    if ctx["blockers"]:
        for item in ctx["blockers"][:10]:
            doc.add_paragraph(f"{item['title']}: {item['summary']}", style="List Bullet")
    else:
        doc.add_paragraph("Aucun blocage sur la periode.", style="List Bullet")

    doc.add_heading("3) Priorites 7 prochains jours", level=1)
    if ctx["next_priorities"]:
        for item in ctx["next_priorities"][:10]:
            doc.add_paragraph(
                f"{item['title']} | {item['next_action']} | Echeance: {item['next_action_date']} | Risque: {item['risk_level']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Aucune priorite planifiee.", style="List Bullet")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    filename = f"rapport_hebdo_{ctx['week_start_s']}_au_{ctx['week_end_s']}.docx"
    log_action("export", "weekly_report", "docx", {"from": ctx["week_start_s"], "to": ctx["week_end_s"]})
    return Response(
        stream.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

